import py7zr
from docx import Document

def extract_file_names_from_7z(archive_path):
    # 使用 py7zr 库读取 7z 文件中的内容
    with py7zr.SevenZipFile(archive_path, mode='r') as archive:
        file_names = archive.getnames()  # 获取所有文件名
    return file_names

def write_to_docx(file_names, output_doc_path):
    # 创建一个 Word 文档
    doc = Document()
    doc.add_heading('7z Archive File List', 0)
    
    # 向文档中添加文件名
    for file_name in file_names:
        doc.add_paragraph(file_name)
    
    # 保存文档
    doc.save(output_doc_path)

if __name__ == "__main__":
    archive_path = 'Downloads.7z'  # 这里是你的 7z 文件路径
    output_doc_path = 'output.docx'  # 输出的 Word 文档路径
    
    # 获取 7z 文件中的所有文件名
    file_names = extract_file_names_from_7z(archive_path)
    
    # 将文件名写入 Word 文档
    write_to_docx(file_names, output_doc_path)
    
    print(f"文件名已成功写入 {output_doc_path}")
