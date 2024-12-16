import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import os
import json

# 隐藏Streamlit默认样式
st.set_page_config(page_title="LIMS Test Report Generator", layout="wide", initial_sidebar_state="auto")
st.markdown(
    """
    <style>
    #MainMenu {visibility:hidden;}
    footer {visibility:hidden;}
    header {visibility:hidden;}
    </style>
    """, 
    unsafe_allow_html=True
)

# 模拟从 LIMS 系统获取数据的函数
def get_lims_test_data():
    # 模拟的 LIMS 测试数据，实际应用中可以通过 API 或数据库获取
    data = {
        'Sample ID': ['S1', 'S2', 'S3', 'S4'],
        'Test Result': [95, 88, 92, 77],
        'Test Date': ['2024-01-01', '2024-01-02', '2024-01-03', '2024-01-04'],
        'Test Type': ['Blood', 'Urine', 'Blood', 'Urine']
    }
    return pd.DataFrame(data)

# 渲染测试数据表格
def render_test_data(df):
    st.subheader("Test Data")
    st.write(df)

# 渲染数据的可视化图表
def render_test_data_visualization(df):
    st.subheader("Test Results Visualization")
    fig = px.bar(df, x='Sample ID', y='Test Result', color='Test Type', title="Test Results by Sample")
    st.plotly_chart(fig)

# 生成并下载报告（Word 格式）
def generate_report(df):
    # 创建 Word 文档
    doc = Document()
    doc.add_heading('LIMS Test Report', 0)

    # 添加数据表格
    doc.add_heading('Test Data', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    # 填充数据
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # 添加图表截图（可以生成图表并嵌入图片）
    doc.add_paragraph("\nVisualized Test Results (bar chart):")
    # 示例：插入图片（假设已经将图表保存为文件）
    # fig.write_image("test_results.png")
    # doc.add_picture('test_results.png', width=Inches(4.0))

    # 保存报告到临时文件
    report_path = "lims_test_report.docx"
    doc.save(report_path)
    return report_path

# 主应用部分
def main():
    # 获取 LIMS 测试数据
    df = get_lims_test_data()

    # 显示测试数据
    render_test_data(df)

    # 渲染数据的可视化
    render_test_data_visualization(df)

    # 生成报告并提供下载
    st.subheader("Download Test Report")
    report_path = generate_report(df)
    with open(report_path, "rb") as f:
        st.download_button("Download Report", f, file_name="lims_test_report.docx")

# 执行主函数
if __name__ == "__main__":
    main()
